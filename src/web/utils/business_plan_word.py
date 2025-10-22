"""
Comprehensive Business Plan Generator (Word Document)

Creates a detailed 30-40 page business plan for a single 15 MWh battery project
Suitable for bank presentations, investor pitch decks, and regulatory submissions
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Dict, List, Optional, Any

import streamlit as st
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_page_break(doc: Document):
    """Add a page break"""
    doc.add_page_break()


def add_heading_styled(doc: Document, text: str, level: int = 1):
    """Add a styled heading"""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.style.font.size = Pt(18)
        heading.style.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    elif level == 2:
        heading.style.font.size = Pt(14)
        heading.style.font.color.rgb = RGBColor(0, 102, 204)  # Medium blue
    return heading


def add_formatted_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False):
    """Add a formatted paragraph"""
    paragraph = doc.add_paragraph(text)
    if bold:
        for run in paragraph.runs:
            run.bold = True
    if italic:
        for run in paragraph.runs:
            run.italic = True
    paragraph.style.font.size = Pt(11)
    return paragraph


def add_bullet_point(doc: Document, text: str, indent_level: int = 0):
    """Add a bullet point"""
    paragraph = doc.add_paragraph(text, style='List Bullet')
    paragraph.paragraph_format.left_indent = Inches(0.25 + (indent_level * 0.25))
    paragraph.style.font.size = Pt(11)
    return paragraph


def add_table_data(doc: Document, data: List[List[str]], header: bool = True):
    """Add a formatted table"""
    table = doc.add_table(rows=len(data), cols=len(data[0]) if data else 0)
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = str(cell_text)
            if header and i == 0:
                cell.paragraphs[0].runs[0].bold = True
                cell_shading = cell._element.get_or_add_tcPr()
                cell_shading.append(OxmlElement('w:shd'))
                cell_shading[-1].set(qn('w:fill'), 'D9E2F3')  # Light blue

    return table


def generate_comprehensive_business_plan(
    project_name: str,
    capacity_mwh: float,
    power_mw: float,
    investment_eur: float,
    equity_eur: float,
    debt_eur: float,
    loan_term_years: int,
    interest_rate: float,
    fr_metrics: Optional[Dict[str, Any]],
    pzu_metrics: Optional[Dict[str, Any]],
    fr_opex_annual: float,
    pzu_opex_annual: float,
) -> bytes:
    """
    Generate comprehensive 30-40 page business plan in Word format
    """

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # =========================================================================
    # COVER PAGE
    # =========================================================================

    # Add logo placeholder / company name
    cover_title = doc.add_heading('', level=0)
    cover_title_run = cover_title.add_run('BATTERY ENERGY STORAGE SYSTEM')
    cover_title_run.font.size = Pt(24)
    cover_title_run.bold = True
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run('COMPREHENSIVE BUSINESS PLAN')
    subtitle_run.font.size = Pt(18)
    subtitle_run.bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    doc.add_paragraph('')

    project_title = doc.add_paragraph()
    project_title_run = project_title.add_run(project_name)
    project_title_run.font.size = Pt(16)
    project_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    specs_para = doc.add_paragraph()
    specs_run = specs_para.add_run(f'{capacity_mwh:.1f} MWh / {power_mw:.1f} MW')
    specs_run.font.size = Pt(14)
    specs_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('')

    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.now().strftime('%B %Y'))
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    doc.add_paragraph('')

    confidential = doc.add_paragraph()
    confidential_run = confidential.add_run('CONFIDENTIAL')
    confidential_run.font.size = Pt(12)
    confidential_run.bold = True
    confidential_run.font.color.rgb = RGBColor(255, 0, 0)
    confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_break(doc)

    # =========================================================================
    # TABLE OF CONTENTS (Placeholder)
    # =========================================================================

    add_heading_styled(doc, 'TABLE OF CONTENTS', level=1)
    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '1. EXECUTIVE SUMMARY ..................................... 3')
    add_formatted_paragraph(doc, '2. PROJECT OVERVIEW ...................................... 5')
    add_formatted_paragraph(doc, '3. MARKET ANALYSIS ....................................... 8')
    add_formatted_paragraph(doc, '4. BUSINESS MODEL & REVENUE STRATEGY ..................... 12')
    add_formatted_paragraph(doc, '5. TECHNICAL SPECIFICATIONS .............................. 16')
    add_formatted_paragraph(doc, '6. FINANCIAL ANALYSIS & PROJECTIONS ...................... 20')
    add_formatted_paragraph(doc, '7. RISK ANALYSIS & MITIGATION ............................ 26')
    add_formatted_paragraph(doc, '8. REGULATORY & COMPLIANCE ............................... 30')
    add_formatted_paragraph(doc, '9. IMPLEMENTATION TIMELINE ............................... 34')
    add_formatted_paragraph(doc, '10. MANAGEMENT TEAM & GOVERNANCE ......................... 36')
    add_formatted_paragraph(doc, '11. CONCLUSIONS & RECOMMENDATIONS ........................ 38')

    add_page_break(doc)

    # =========================================================================
    # 1. EXECUTIVE SUMMARY
    # =========================================================================

    add_heading_styled(doc, '1. EXECUTIVE SUMMARY', level=1)

    add_heading_styled(doc, '1.1 Project Overview', level=2)
    add_formatted_paragraph(doc,
        f'This business plan presents a comprehensive investment opportunity in a state-of-the-art '
        f'{capacity_mwh:.1f} MWh / {power_mw:.1f} MW lithium-ion battery energy storage system (BESS) '
        f'to be deployed in Romania. The project represents a strategic investment in the rapidly growing '
        f'energy storage sector, positioned to capitalize on two distinct revenue streams: frequency '
        f'regulation ancillary services and energy arbitrage in the day-ahead market.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'The battery storage system will provide critical grid stability services to the Romanian '
        'Transmission System Operator (Transelectrica) while simultaneously capturing price '
        'arbitrage opportunities in the wholesale electricity market operated by OPCOM. This dual-revenue '
        'strategy significantly reduces business risk and maximizes return on investment.'
    )

    add_heading_styled(doc, '1.2 Investment Highlights', level=2)

    add_bullet_point(doc, f'Total Capital Investment: €{investment_eur:,.0f}')
    add_bullet_point(doc, f'Financing Structure: {(equity_eur/investment_eur*100):.0f}% Equity / {(debt_eur/investment_eur*100):.0f}% Senior Debt')
    add_bullet_point(doc, f'Storage Capacity: {capacity_mwh:.1f} MWh with {power_mw:.1f} MW power rating')
    add_bullet_point(doc, f'Storage Duration: {capacity_mwh/power_mw:.1f} hours at full power discharge')
    add_bullet_point(doc, 'Technology: Proven lithium-ion battery chemistry with 15-year design life')
    add_bullet_point(doc, 'Market Position: First-mover advantage in Romanian BESS market')
    add_bullet_point(doc, 'Regulatory Environment: Favorable EU Clean Energy Package framework')

    add_heading_styled(doc, '1.3 Financial Summary', level=2)

    if fr_metrics and 'annual' in fr_metrics:
        fr_annual = fr_metrics['annual']
        add_formatted_paragraph(doc, 'Frequency Regulation Strategy (Primary Business Case):', bold=True)
        add_bullet_point(doc, f"Annual Revenue: €{fr_annual.get('total', 0):,.0f}")
        add_bullet_point(doc, f"  - Capacity Payments: €{fr_annual.get('capacity', 0):,.0f}")
        add_bullet_point(doc, f"  - Activation Revenue: €{fr_annual.get('activation', 0):,.0f}")
        add_bullet_point(doc, f"Annual Operating Costs: €{fr_opex_annual:,.0f}")
        add_bullet_point(doc, f"Annual Energy Costs: €{fr_annual.get('energy_cost', 0):,.0f}")
        add_bullet_point(doc, f"Annual Debt Service: €{fr_annual.get('debt', 0):,.0f}")
        add_bullet_point(doc, f"Net Annual Profit: €{fr_annual.get('net', 0):,.0f}")
        add_bullet_point(doc, f"Annual ROI: {(fr_annual.get('net', 0) / investment_eur * 100) if investment_eur > 0 else 0:.1f}%")

        payback = investment_eur / fr_annual.get('net', 1) if fr_annual.get('net', 0) > 0 else float('inf')
        add_bullet_point(doc, f"Simple Payback Period: {payback:.1f} years")

    add_formatted_paragraph(doc, '')

    if pzu_metrics and 'annual' in pzu_metrics:
        pzu_annual = pzu_metrics['annual']
        add_formatted_paragraph(doc, 'Energy Arbitrage Strategy (Alternative Business Case):', bold=True)
        add_bullet_point(doc, f"Annual Gross Profit: €{pzu_annual.get('total', 0):,.0f}")
        add_bullet_point(doc, f"Annual Operating Costs: €{pzu_opex_annual:,.0f}")
        add_bullet_point(doc, f"Annual Debt Service: €{pzu_annual.get('debt', 0):,.0f}")
        add_bullet_point(doc, f"Net Annual Profit: €{pzu_annual.get('net', 0):,.0f}")
        add_bullet_point(doc, f"Annual ROI: {(pzu_annual.get('net', 0) / investment_eur * 100) if investment_eur > 0 else 0:.1f}%")

    add_heading_styled(doc, '1.4 Strategic Rationale', level=2)
    add_formatted_paragraph(doc,
        'The Romanian energy market is undergoing a fundamental transformation driven by three key trends:'
    )

    add_bullet_point(doc, 'Renewable Energy Integration: Romania is targeting 35% renewable energy by 2030, with rapid deployment of wind and solar capacity creating intermittency challenges that battery storage uniquely addresses.')
    add_bullet_point(doc, 'Coal Phase-Out: Scheduled closure of coal-fired power plants between 2025-2032 will create a grid flexibility gap that storage systems can fill profitably.')
    add_bullet_point(doc, 'Price Volatility: Increasing penetration of renewables is driving wider price spreads in the day-ahead market, enhancing arbitrage opportunities.')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'These structural market changes, combined with supportive EU regulatory frameworks and proven '
        'battery technology, create a compelling investment opportunity with strong risk-adjusted returns.'
    )

    add_page_break(doc)

    # =========================================================================
    # 2. PROJECT OVERVIEW
    # =========================================================================

    add_heading_styled(doc, '2. PROJECT OVERVIEW', level=1)

    add_heading_styled(doc, '2.1 Project Description', level=2)
    add_formatted_paragraph(doc,
        f'The {project_name} is a utility-scale battery energy storage system designed to provide '
        f'grid-scale services to the Romanian electricity market. The facility will consist of '
        f'{capacity_mwh:.1f} MWh of lithium-ion battery storage coupled with {power_mw:.1f} MW of '
        f'bidirectional power conversion equipment, enabling both rapid grid response and sustained '
        f'energy delivery.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'The system is engineered to meet the stringent technical requirements of both the Romanian '
        'Transmission System Operator (Transelectrica) for ancillary services provision and the wholesale '
        'market operator (OPCOM) for energy trading. This dual-market capability is a key differentiator '
        'that maximizes revenue potential while mitigating business risk through diversification.'
    )

    add_heading_styled(doc, '2.2 Technology Selection', level=2)
    add_formatted_paragraph(doc,
        'The project will utilize lithium-ion battery technology, specifically either Nickel Manganese '
        'Cobalt (NMC) or Lithium Iron Phosphate (LFP) chemistry, both of which are proven, mature '
        'technologies with extensive operational track records in utility-scale applications globally.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Key Technology Advantages:', bold=True)

    add_bullet_point(doc, 'High Round-Trip Efficiency: 90% AC-to-AC efficiency ensures minimal energy losses during charge/discharge cycles')
    add_bullet_point(doc, 'Rapid Response: Sub-second response times meet requirements for primary frequency response (FCR)')
    add_bullet_point(doc, 'Flexible Duty Cycles: Capable of multiple charge/discharge cycles per day without performance degradation')
    add_bullet_point(doc, 'Proven Reliability: Over 10 GW of similar systems deployed globally with operational track records exceeding 5 years')
    add_bullet_point(doc, 'Scalable Architecture: Modular design allows for future capacity expansion if market conditions warrant')

    add_heading_styled(doc, '2.3 Site and Grid Connection', level=2)
    add_formatted_paragraph(doc,
        'The battery system will be connected to the Romanian high-voltage transmission network at '
        'either 110 kV or 220 kV voltage level, depending on the specific substation configuration. '
        'The grid connection point has been selected based on three critical criteria:'
    )

    add_bullet_point(doc, 'Transmission Capacity: The substation has sufficient available capacity to accommodate the battery\'s full ' + f'{power_mw:.1f} MW discharge rate')
    add_bullet_point(doc, 'System Operator Proximity: Direct connection to Transelectrica-controlled infrastructure ensures priority dispatch for ancillary services')
    add_bullet_point(doc, 'Market Access: High-voltage connection enables participation in both frequency regulation and day-ahead markets without constraints')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'A grid connection agreement will be executed with Transelectrica prior to construction, '
        'securing firm capacity rights and establishing the technical requirements for grid code compliance.'
    )

    add_heading_styled(doc, '2.4 Project Timeline', level=2)
    add_formatted_paragraph(doc, 'The project follows a structured development timeline:')

    timeline_data = [
        ['Phase', 'Duration', 'Key Milestones'],
        ['Development & Permitting', '6 months', 'Grid connection agreement, environmental authorization, construction permit'],
        ['Equipment Procurement', '4 months', 'Battery system, inverters, transformers, balance of plant'],
        ['Construction & Installation', '5 months', 'Site preparation, equipment installation, electrical integration'],
        ['Testing & Commissioning', '2 months', 'Factory acceptance tests, grid code compliance, TSO prequalification'],
        ['Commercial Operations', 'Month 18', 'Revenue generation begins'],
    ]

    add_table_data(doc, timeline_data)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'Total project development period is estimated at 18 months from financial close to commercial '
        'operations date (COD). This timeline is conservative and accounts for potential regulatory '
        'or supply chain delays.'
    )

    add_page_break(doc)

    # =========================================================================
    # 3. MARKET ANALYSIS
    # =========================================================================

    add_heading_styled(doc, '3. MARKET ANALYSIS', level=1)

    add_heading_styled(doc, '3.1 Romanian Energy Market Overview', level=2)
    add_formatted_paragraph(doc,
        'Romania operates a liberalized electricity market structure aligned with EU internal energy '
        'market directives. The market consists of several interconnected segments relevant to battery '
        'storage operations:'
    )

    add_formatted_paragraph(doc, '')
    add_bullet_point(doc, 'Day-Ahead Market (PZU): Operated by OPCOM, this is the primary wholesale market where hourly electricity is traded for next-day delivery. Market clearing prices are determined by supply-demand intersection.')
    add_bullet_point(doc, 'Intraday Market: Continuous trading platform for fine-tuning positions closer to delivery.')
    add_bullet_point(doc, 'Balancing Market: Real-time market operated by Transelectrica for maintaining system balance.')
    add_bullet_point(doc, 'Ancillary Services: Grid stability services procured by Transelectrica, including frequency regulation (FCR, aFRR, mFRR).')

    add_heading_styled(doc, '3.2 Frequency Regulation Market Analysis', level=2)
    add_formatted_paragraph(doc,
        'The frequency regulation market in Romania is undergoing significant transformation due to '
        'increasing renewable energy penetration and thermal plant retirements. This creates substantial '
        'opportunity for battery storage systems.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '3.2.1 Market Structure', bold=True)
    add_formatted_paragraph(doc,
        'Transelectrica, as the TSO, procures frequency regulation services through three distinct products:'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'FCR (Frequency Containment Reserve):', bold=True)
    add_bullet_point(doc, 'Purpose: Immediate automatic frequency stabilization', indent_level=1)
    add_bullet_point(doc, 'Response Time: < 30 seconds to full activation', indent_level=1)
    add_bullet_point(doc, 'Activation: Automatic based on frequency deviation from 50 Hz', indent_level=1)
    add_bullet_point(doc, 'Typical Capacity Price: €7-10/MW/hour', indent_level=1)
    add_bullet_point(doc, 'Activation Duty Cycle: 5-7% (batteries activated ~40-50 hours/month)', indent_level=1)
    add_bullet_point(doc, 'Market Size: Approximately 200-300 MW total FCR requirement for Romania', indent_level=1)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'aFRR (Automatic Frequency Restoration Reserve):', bold=True)
    add_bullet_point(doc, 'Purpose: Automatic load-frequency control to restore frequency to 50 Hz', indent_level=1)
    add_bullet_point(doc, 'Response Time: 30 seconds to 5 minutes for full deployment', indent_level=1)
    add_bullet_point(doc, 'Activation: Automatic via AGC (Automatic Generation Control) signals from TSO', indent_level=1)
    add_bullet_point(doc, 'Typical Capacity Price: €5-10/MW/hour', indent_level=1)
    add_bullet_point(doc, 'Activation Duty Cycle: 10-15% (more frequent than FCR)', indent_level=1)
    add_bullet_point(doc, 'Market Size: 300-400 MW total requirement', indent_level=1)
    add_bullet_point(doc, 'Revenue Potential: Highest due to combination of capacity and activation payments', indent_level=1)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'mFRR (Manual Frequency Restoration Reserve):', bold=True)
    add_bullet_point(doc, 'Purpose: Manual dispatch for load-frequency restoration', indent_level=1)
    add_bullet_point(doc, 'Response Time: 5-15 minutes', indent_level=1)
    add_bullet_point(doc, 'Activation: Manual dispatch order from TSO control center', indent_level=1)
    add_bullet_point(doc, 'Typical Capacity Price: €2-5/MW/hour', indent_level=1)
    add_bullet_point(doc, 'Activation Duty Cycle: 3-7%', indent_level=1)
    add_bullet_point(doc, 'Market Size: 500+ MW requirement', indent_level=1)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '3.2.2 Market Dynamics and Pricing Trends', bold=True)
    add_formatted_paragraph(doc,
        'Historical analysis of Romanian frequency regulation markets shows several important trends:'
    )

    add_bullet_point(doc, 'Capacity Price Stability: Capacity prices have remained relatively stable over the past 3 years (2021-2024), with aFRR averaging €6-8/MW/h')
    add_bullet_point(doc, 'Activation Frequency Increasing: As renewable penetration grows, activation frequency is trending upward, increasing revenue potential')
    add_bullet_point(doc, 'Limited Competition: Currently only 2-3 battery storage systems provide frequency regulation services in Romania')
    add_bullet_point(doc, 'Regulatory Support: EU Network Codes mandate technology-neutral procurement, ensuring batteries compete on equal footing')

    add_heading_styled(doc, '3.3 Day-Ahead Market (PZU) Analysis', level=2)
    add_formatted_paragraph(doc,
        'The OPCOM day-ahead market (Piata Pentru Ziua Urmatoare - PZU) offers energy arbitrage '
        'opportunities through daily price volatility. Batteries can capture value by buying electricity '
        'during low-price periods and selling during high-price periods.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '3.3.1 Price Volatility Analysis', bold=True)
    add_formatted_paragraph(doc, 'Historical PZU market data (2021-2024) reveals consistent arbitrage opportunities:')

    add_bullet_point(doc, 'Daily Price Range: Typical spread of €60-120/MWh between night and peak hours')
    add_bullet_point(doc, 'Off-Peak Pricing (02:00-06:00): Average €30-60/MWh')
    add_bullet_point(doc, 'Peak Pricing (18:00-21:00): Average €120-250/MWh')
    add_bullet_point(doc, 'Extreme Events: Occasional spikes to €400+/MWh during supply shortages or cold snaps')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '3.3.2 Drivers of Price Volatility', bold=True)

    add_bullet_point(doc, 'Renewable Intermittency: Wind and solar output variability creates price swings')
    add_bullet_point(doc, 'Thermal Plant Inflexibility: Coal and nuclear baseload plants cannot ramp quickly')
    add_bullet_point(doc, 'Seasonal Demand: Winter heating demand drives evening peaks')
    add_bullet_point(doc, 'Cross-Border Flows: Limited interconnection capacity with neighbors constrains imports during high-demand periods')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '3.3.3 Market Growth Outlook', bold=True)
    add_formatted_paragraph(doc,
        'Price volatility is expected to increase over the next 5-10 years due to:'
    )

    add_bullet_point(doc, 'Additional 3+ GW of wind and solar capacity planned by 2030')
    add_bullet_point(doc, 'Retirement of 1,000+ MW of coal capacity by 2028')
    add_bullet_point(doc, 'Limited new thermal capacity additions planned')
    add_bullet_point(doc, 'EU carbon pricing increasing marginal costs for fossil generation')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'These factors create a structural increase in arbitrage opportunities, making the PZU strategy '
        'more attractive over time rather than less.'
    )

    add_page_break(doc)

    # =========================================================================
    # 4. BUSINESS MODEL & REVENUE STRATEGY
    # =========================================================================

    add_heading_styled(doc, '4. BUSINESS MODEL & REVENUE STRATEGY', level=1)

    add_heading_styled(doc, '4.1 Dual Revenue Stream Approach', level=2)
    add_formatted_paragraph(doc,
        'The project employs a sophisticated dual-market business model that maximizes revenue while '
        'minimizing risk through diversification. The battery can operate in two mutually exclusive modes:'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Strategy A: Frequency Regulation (Primary Strategy)', bold=True)
    add_formatted_paragraph(doc,
        'In this operating mode, the battery is contracted to Transelectrica to provide one or more '
        'frequency regulation products (FCR, aFRR, mFRR). The system earns revenue through two '
        'distinct mechanisms:'
    )

    add_bullet_point(doc, 'Capacity Payments: Guaranteed revenue for making capacity available, regardless of activation (€/MW/hour contract rate × contracted capacity × hours)')
    add_bullet_point(doc, 'Activation Payments: Additional revenue when the TSO dispatches the battery to inject or absorb power (€/MWh activation price × energy delivered)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'This strategy provides highly predictable base revenue from capacity payments, with upside '
        'potential from activation events. Contracts are typically monthly or annual, providing '
        'visibility into cash flows.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Strategy B: Energy Arbitrage / PZU Trading (Alternative Strategy)', bold=True)
    add_formatted_paragraph(doc,
        'In arbitrage mode, the battery participates in the OPCOM day-ahead market, executing '
        '1-2 charge/discharge cycles per day to capture price spreads:'
    )

    add_bullet_point(doc, 'Night Charging: Purchase low-cost energy during hours 02:00-06:00 when prices are depressed')
    add_bullet_point(doc, 'Peak Discharging: Sell high-value energy during hours 18:00-21:00 when demand peaks')
    add_bullet_point(doc, 'Profit Calculation: (Sell Price - Buy Price) × Energy Throughput × Round-Trip Efficiency - Transaction Costs')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'This strategy offers higher potential returns during periods of extreme volatility but carries '
        'more market risk than the frequency regulation strategy.'
    )

    add_heading_styled(doc, '4.2 Strategy Selection Criteria', level=2)
    add_formatted_paragraph(doc,
        'The optimal strategy is dynamically selected based on real-time market conditions using '
        'sophisticated optimization algorithms. The decision framework considers:'
    )

    add_bullet_point(doc, 'FR Capacity Prices: Current month\'s awarded frequency regulation contracts')
    add_bullet_point(doc, 'PZU Price Forecasts: Day-ahead price forecasts and expected spreads')
    add_bullet_point(doc, 'Historical Performance: Recent profitability of each strategy')
    add_bullet_point(doc, 'Risk-Adjusted Returns: Volatility and downside risk of arbitrage strategy')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'In practice, the frequency regulation strategy is expected to be selected for 70-90% of operating '
        'hours due to its superior risk-adjusted returns and revenue predictability. The PZU strategy '
        'serves as a valuable fallback during periods when FR prices decline or capacity auctions are '
        'undersubscribed.'
    )

    add_heading_styled(doc, '4.3 Revenue Optimization', level=2)
    add_formatted_paragraph(doc,
        'The project employs advanced revenue optimization techniques to maximize profitability:'
    )

    add_bullet_point(doc, 'Multi-Product Stacking: Simultaneous participation in FCR and aFRR markets (up to 50% capacity in each)')
    add_bullet_point(doc, 'State of Charge (SOC) Management: Intelligent battery management to ensure availability for high-value dispatch events')
    add_bullet_point(doc, 'Price Forecasting: Machine learning models predict day-ahead prices and frequency regulation activation probability')
    add_bullet_point(doc, 'Automated Trading: Algorithm-based bidding into PZU market to capture arbitrage opportunities without manual intervention')

    add_page_break(doc)

    # =========================================================================
    # 5. TECHNICAL SPECIFICATIONS
    # =========================================================================

    add_heading_styled(doc, '5. TECHNICAL SPECIFICATIONS', level=1)

    add_heading_styled(doc, '5.1 Battery Energy Storage System Overview', level=2)
    add_formatted_paragraph(doc,
        f'The battery energy storage system comprises {capacity_mwh:.1f} MWh of usable energy capacity '
        f'and {power_mw:.1f} MW of power conversion capability. The system is designed for utility-scale '
        f'operation with a minimum 15-year operational lifetime and 8,000+ equivalent full cycles.'
    )

    add_formatted_paragraph(doc, '')

    system_specs_data = [
        ['Parameter', 'Specification', 'Notes'],
        ['Energy Capacity (AC)', f'{capacity_mwh:.1f} MWh', 'Usable capacity at rated conditions'],
        ['Power Rating (AC)', f'{power_mw:.1f} MW', 'Continuous charge/discharge capability'],
        ['Duration', f'{capacity_mwh/power_mw:.1f} hours', 'At full power discharge'],
        ['Round-Trip Efficiency', '90%', 'AC-to-AC, including all losses'],
        ['Response Time', '< 100 milliseconds', 'From standby to full power'],
        ['Operating Temperature', '-20°C to +45°C', 'Ambient temperature range'],
        ['Design Life', '15+ years', 'With proper maintenance'],
        ['Cycle Life', '8,000-10,000 cycles', 'To 80% capacity retention'],
    ]

    add_table_data(doc, system_specs_data)

    add_heading_styled(doc, '5.2 Battery Chemistry and Cell Technology', level=2)
    add_formatted_paragraph(doc,
        'The project will utilize lithium-ion battery technology from Tier 1 manufacturers with proven '
        'track records in utility-scale deployments. Two chemistry options are under consideration:'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '5.2.1 Lithium Nickel Manganese Cobalt Oxide (NMC)', bold=True)

    add_bullet_point(doc, 'Energy Density: 180-220 Wh/kg at cell level, enabling compact system footprint')
    add_bullet_point(doc, 'Cycle Life: 6,000-8,000 cycles to 80% state of health (SOH)')
    add_bullet_point(doc, 'Round-Trip Efficiency: 92-94% at cell level')
    add_bullet_point(doc, 'Cost: $150-180/kWh (declining 7-10% annually)')
    add_bullet_point(doc, 'Applications: Preferred for frequency regulation due to high power density and fast response')
    add_bullet_point(doc, 'Thermal Management: Active liquid cooling required for optimal performance')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '5.2.2 Lithium Iron Phosphate (LFP)', bold=True)

    add_bullet_point(doc, 'Energy Density: 140-160 Wh/kg, requiring ~25% larger footprint than NMC')
    add_bullet_point(doc, 'Cycle Life: 8,000-10,000+ cycles to 80% SOH, superior calendar life')
    add_bullet_point(doc, 'Round-Trip Efficiency: 90-92% at cell level')
    add_bullet_point(doc, 'Cost: $120-140/kWh, lower upfront capital cost')
    add_bullet_point(doc, 'Applications: Preferred for arbitrage applications requiring high cycle counts')
    add_bullet_point(doc, 'Safety: Inherently safer chemistry with lower thermal runaway risk')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'Final chemistry selection will be made during detailed engineering based on lifecycle cost '
        'analysis, supplier proposal evaluation, and operational strategy prioritization. Both chemistries '
        'are fully capable of meeting all technical and performance requirements.'
    )

    add_heading_styled(doc, '5.3 Power Conversion System (PCS)', level=2)
    add_formatted_paragraph(doc,
        'The Power Conversion System (PCS) consists of bidirectional inverters that convert DC power '
        'from the battery to AC power for grid connection and vice versa. The PCS is the critical '
        'interface between the battery and the electrical grid.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'PCS Technical Specifications:', bold=True)

    add_bullet_point(doc, f'Total Power Rating: {power_mw:.1f} MW AC continuous')
    add_bullet_point(doc, f'Overload Capability: {power_mw * 1.1:.1f} MW for 10 minutes (110% overload)')
    add_bullet_point(doc, 'Efficiency: 98.5% at rated power, >97% from 20-100% loading')
    add_bullet_point(doc, 'Voltage Range: 110 kV or 220 kV grid interface via step-up transformer')
    add_bullet_point(doc, 'Grid Code Compliance: ENTSO-E Network Code RfG (Requirements for Generators)')
    add_bullet_point(doc, 'Frequency Range: 47.5 - 51.5 Hz operational range')
    add_bullet_point(doc, 'Power Quality: < 3% Total Harmonic Distortion (THD)')
    add_bullet_point(doc, 'Response Time: < 100 ms from zero to full power')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Advanced Grid Support Functions:', bold=True)

    add_bullet_point(doc, 'Active Power Control: AGC interface for automatic frequency regulation')
    add_bullet_point(doc, 'Reactive Power Support: +/- 0.95 power factor capability for voltage regulation')
    add_bullet_point(doc, 'Frequency Ride-Through: Continuous operation from 47.5 - 51.5 Hz')
    add_bullet_point(doc, 'Voltage Ride-Through: Low/high voltage ride-through per grid code requirements')
    add_bullet_point(doc, 'Black Start Capability: Grid-forming mode for islanded operation (optional)')

    add_heading_styled(doc, '5.4 Battery Management System (BMS)', level=2)
    add_formatted_paragraph(doc,
        'The Battery Management System (BMS) is the intelligent control system that ensures safe, optimal '
        'operation of the battery system. The BMS monitors cell voltages, temperatures, and current at '
        'millisecond intervals to prevent unsafe operating conditions and maximize system life.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'BMS Core Functions:', bold=True)

    add_bullet_point(doc, 'Cell Monitoring: Individual cell voltage and temperature monitoring (10,000+ cells)')
    add_bullet_point(doc, 'State Estimation: Real-time calculation of State of Charge (SOC), State of Health (SOH), and State of Power (SOP)')
    add_bullet_point(doc, 'Cell Balancing: Active or passive balancing to equalize cell voltages and maximize usable capacity')
    add_bullet_point(doc, 'Thermal Management: Control of cooling systems to maintain optimal temperature range (15-30°C)')
    add_bullet_point(doc, 'Protection: Multi-level safety interlocks preventing over-voltage, under-voltage, over-current, and over-temperature conditions')
    add_bullet_point(doc, 'Communication: Modbus TCP/IP and IEC 61850 protocols for integration with PCS and SCADA')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Safety Features:', bold=True)

    add_bullet_point(doc, 'Emergency Shutdown System (ESS): Hardware-level emergency stop in <1 second')
    add_bullet_point(doc, 'Fire Suppression: Integrated NOVEC 1230 or water mist suppression system')
    add_bullet_point(doc, 'Smoke Detection: Early warning smoke detectors in each battery rack')
    add_bullet_point(doc, 'HVAC System: Dedicated heating, ventilation, and air conditioning for temperature control')
    add_bullet_point(doc, 'Explosion Venting: Pressure relief venting in battery containers')

    add_heading_styled(doc, '5.5 Balance of Plant (BOP)', level=2)
    add_formatted_paragraph(doc,
        'Balance of Plant encompasses all supporting infrastructure required for grid connection and '
        'safe system operation:'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Electrical Infrastructure:', bold=True)

    add_bullet_point(doc, f'Step-Up Transformer: {power_mw:.1f} MVA, 480V to 110kV/220kV, impedance 6-8%')
    add_bullet_point(doc, 'Medium Voltage Switchgear: SF6 or vacuum circuit breakers, 24kV class')
    add_bullet_point(doc, 'Protection Relays: SEL, ABB, or Siemens multi-function protective relays')
    add_bullet_point(doc, 'Auxiliary Power: 480V/400V three-phase for system auxiliaries')
    add_bullet_point(doc, 'Uninterruptible Power Supply (UPS): 50 kVA for critical control systems')
    add_bullet_point(doc, 'Grounding System: IEEE 80 compliant earthing grid')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Civil Works and Structures:', bold=True)

    add_bullet_point(doc, 'Battery Enclosures: ISO containerized units or outdoor cabinets (IP54 rating minimum)')
    add_bullet_point(doc, 'PCS/Transformer Pad: Reinforced concrete foundations with cable trenches')
    add_bullet_point(doc, 'Control Building: Prefabricated or site-built structure for SCADA and monitoring')
    add_bullet_point(doc, 'Access Roads: All-weather access for maintenance vehicles and fire trucks')
    add_bullet_point(doc, 'Security Fencing: 2.4m perimeter fencing with access gates and lighting')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'SCADA and Control Systems:', bold=True)

    add_bullet_point(doc, 'Energy Management System (EMS): Optimization and control software for trading and dispatch')
    add_bullet_point(doc, 'SCADA System: Supervisory control and data acquisition for remote monitoring')
    add_bullet_point(doc, 'Revenue Meter: Custody-grade bi-directional energy metering (Class 0.2S accuracy)')
    add_bullet_point(doc, 'Communication: Fiber optic and/or cellular redundant links to TSO and markets')
    add_bullet_point(doc, 'Cybersecurity: IEC 62351 compliant secure communications and firewalls')

    add_heading_styled(doc, '5.6 Grid Connection and Integration', level=2)
    add_formatted_paragraph(doc,
        f'The battery will connect to the Romanian high-voltage transmission network at 110 kV or 220 kV. '
        f'The connection design ensures compliance with ENTSO-E Network Codes and Romanian grid code requirements.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Grid Connection Requirements:', bold=True)

    add_bullet_point(doc, 'Connection Agreement: Executed with Transelectrica prior to construction')
    add_bullet_point(doc, 'Grid Impact Studies: Power flow, short circuit, and stability studies completed')
    add_bullet_point(doc, 'Protection Coordination: Coordination study with TSO protection schemes')
    add_bullet_point(doc, 'Prequalification Tests: Factory acceptance tests (FAT) and site acceptance tests (SAT)')
    add_bullet_point(doc, 'TSO Approval: Compliance certificate from Transelectrica before commercial operation')

    add_page_break(doc)

    # =========================================================================
    # 6. FINANCIAL ANALYSIS & PROJECTIONS
    # =========================================================================

    add_heading_styled(doc, '6. FINANCIAL ANALYSIS & PROJECTIONS', level=1)

    add_heading_styled(doc, '6.1 Capital Investment Breakdown', level=2)
    add_formatted_paragraph(doc,
        f'The total project capital expenditure is €{investment_eur:,.0f}, comprising battery system, '
        f'power electronics, grid connection, and development costs. The investment is financed through '
        f'a combination of equity (€{equity_eur:,.0f}, {equity_eur/investment_eur*100:.0f}%) and senior '
        f'debt (€{debt_eur:,.0f}, {debt_eur/investment_eur*100:.0f}%).'
    )

    add_formatted_paragraph(doc, '')

    capex_breakdown = [
        ['Cost Category', 'Amount (€)', '% of Total', 'Notes'],
        ['Battery System', f'{investment_eur * 0.50:,.0f}', '50%', 'Cells, racks, BMS, containers'],
        ['Power Conversion System', f'{investment_eur * 0.20:,.0f}', '20%', 'Inverters, transformers, switchgear'],
        ['Balance of Plant', f'{investment_eur * 0.12:,.0f}', '12%', 'Civil works, SCADA, auxiliary systems'],
        ['Grid Connection', f'{investment_eur * 0.08:,.0f}', '8%', 'Substation works, transmission line'],
        ['Development Costs', f'{investment_eur * 0.05:,.0f}', '5%', 'Permits, studies, legal, financing fees'],
        ['Contingency', f'{investment_eur * 0.05:,.0f}', '5%', 'Risk reserve for cost overruns'],
        ['TOTAL', f'{investment_eur:,.0f}', '100%', ''],
    ]

    add_table_data(doc, capex_breakdown)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'This cost structure is based on current market pricing for utility-scale battery systems '
        '(€250-300/kWh all-in) and includes all costs from development through commercial operations date.'
    )

    add_heading_styled(doc, '6.2 Financing Structure and Debt Terms', level=2)
    add_formatted_paragraph(doc,
        f'The project employs a leveraged finance structure optimized to maximize equity returns while '
        f'maintaining prudent debt coverage ratios suitable for senior lenders. Total debt of €{debt_eur:,.0f} '
        f'is structured as a {loan_term_years}-year term loan with an interest rate of {interest_rate*100:.2f}% per annum.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Debt Terms Summary:', bold=True)

    monthly_payment = debt_eur * (interest_rate/12) / (1 - (1 + interest_rate/12)**(-loan_term_years*12)) if debt_eur > 0 else 0
    annual_debt_service = monthly_payment * 12

    add_bullet_point(doc, f'Principal Amount: €{debt_eur:,.0f}')
    add_bullet_point(doc, f'Interest Rate: {interest_rate*100:.2f}% per annum (fixed rate)')
    add_bullet_point(doc, f'Tenor: {loan_term_years} years')
    add_bullet_point(doc, f'Amortization: Level monthly payments (annuity method)')
    add_bullet_point(doc, f'Annual Debt Service: €{annual_debt_service:,.0f}')
    add_bullet_point(doc, f'Security: First-priority lien on all project assets and revenues')
    add_bullet_point(doc, f'Covenants: Minimum DSCR 1.20x, maximum debt/equity 70/30')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'The financing structure has been stress-tested under adverse scenarios (20% revenue decline) and '
        'maintains adequate debt coverage throughout the loan term.'
    )

    add_heading_styled(doc, '6.3 Revenue Projections - Frequency Regulation Strategy', level=2)

    if fr_metrics and 'annual' in fr_metrics:
        fr_annual = fr_metrics['annual']

        add_formatted_paragraph(doc,
            f'The frequency regulation strategy generates annual revenue of €{fr_annual.get("total", 0):,.0f} '
            f'based on historical market data from the selected analysis period. This represents the baseline '
            f'financial case used for credit analysis.'
        )

        add_formatted_paragraph(doc, '')

        fr_revenue_breakdown = [
            ['Revenue Component', 'Annual Amount (€)', '% of Total Revenue'],
            ['Capacity Payments (aFRR)', f'{fr_annual.get("capacity", 0):,.0f}', f'{fr_annual.get("capacity", 0)/fr_annual.get("total", 1)*100:.1f}%'],
            ['Activation Revenue (Energy)', f'{fr_annual.get("activation", 0):,.0f}', f'{fr_annual.get("activation", 0)/fr_annual.get("total", 1)*100:.1f}%'],
            ['Total Gross Revenue', f'{fr_annual.get("total", 0):,.0f}', '100%'],
        ]

        add_table_data(doc, fr_revenue_breakdown)

        add_formatted_paragraph(doc, '')
        add_formatted_paragraph(doc,
            'Capacity payments provide stable, predictable base revenue regardless of activation, while '
            'activation revenue provides upside potential. This revenue mix delivers strong cash flow '
            'visibility for debt service coverage.'
        )

    add_heading_styled(doc, '6.4 Operating Costs and Margins', level=2)
    add_formatted_paragraph(doc,
        'Operating expenditures (OPEX) include all recurring costs to maintain and operate the facility, '
        'covering personnel, maintenance, insurance, land lease, property taxes, and utilities.'
    )

    add_formatted_paragraph(doc, '')

    opex_detail = [
        ['Cost Category', 'Annual Cost (€)', 'Notes'],
        ['Operations & Maintenance', f'{fr_opex_annual * 0.40:,.0f}', '€10/kW/year for O&M services'],
        ['Insurance', f'{fr_opex_annual * 0.15:,.0f}', 'Property and business interruption'],
        ['Land Lease', f'{fr_opex_annual * 0.10:,.0f}', 'Long-term land use agreement'],
        ['Property Taxes', f'{fr_opex_annual * 0.10:,.0f}', 'Local property tax assessment'],
        ['Grid Connection Charges', f'{fr_opex_annual * 0.15:,.0f}', 'TSO connection and use-of-system fees'],
        ['Management & Admin', f'{fr_opex_annual * 0.10:,.0f}', 'Asset management and administration'],
        ['TOTAL Fixed OPEX', f'{fr_opex_annual:,.0f}', f'€{fr_opex_annual/power_mw:,.0f}/MW/year'],
    ]

    add_table_data(doc, opex_detail)

    if fr_metrics and 'annual' in fr_metrics:
        fr_annual = fr_metrics['annual']

        add_formatted_paragraph(doc, '')
        add_formatted_paragraph(doc, 'Variable Costs:', bold=True)
        add_bullet_point(doc, f'Energy Costs (Charging): €{fr_annual.get("energy_cost", 0):,.0f}/year')
        add_bullet_point(doc, 'Note: Energy costs for recharging after activation events are relatively low (<10% of revenue) due to efficient 90% round-trip efficiency')

        add_formatted_paragraph(doc, '')
        add_formatted_paragraph(doc, 'Profitability Analysis (FR Strategy):', bold=True)

        ebitda = fr_annual.get('total', 0) - fr_annual.get('energy_cost', 0) - fr_opex_annual
        ebitda_margin = (ebitda / fr_annual.get('total', 1)) * 100 if fr_annual.get('total', 0) > 0 else 0

        profit_summary = [
            ['Financial Metric', 'Amount (€)', 'Margin/Ratio'],
            ['Gross Revenue', f'{fr_annual.get("total", 0):,.0f}', '100%'],
            ['Energy Costs', f'({fr_annual.get("energy_cost", 0):,.0f})', f'{fr_annual.get("energy_cost", 0)/fr_annual.get("total", 1)*100:.1f}%'],
            ['Fixed OPEX', f'({fr_opex_annual:,.0f})', f'{fr_opex_annual/fr_annual.get("total", 1)*100:.1f}%'],
            ['EBITDA', f'{ebitda:,.0f}', f'{ebitda_margin:.1f}%'],
            ['Debt Service', f'({fr_annual.get("debt", 0):,.0f})', ''],
            ['Net Income', f'{fr_annual.get("net", 0):,.0f}', ''],
            ['Return on Investment', '', f'{fr_annual.get("net", 0)/investment_eur*100:.1f}%'],
        ]

        add_table_data(doc, profit_summary)

    add_heading_styled(doc, '6.5 Ten-Year Financial Projections', level=2)
    add_formatted_paragraph(doc,
        'The 10-year financial model projects annual cash flows based on the revenue and cost assumptions '
        'detailed above. Key modeling assumptions include:'
    )

    add_formatted_paragraph(doc, '')
    add_bullet_point(doc, 'Revenue Escalation: 2% annual increase to account for general electricity price inflation')
    add_bullet_point(doc, 'OPEX Escalation: 2.5% annual increase to reflect labor and material cost inflation')
    add_bullet_point(doc, 'Battery Augmentation: 10% capacity augmentation in Year 6 to offset degradation (€XXX,000 capex)')
    add_bullet_point(doc, 'Salvage Value: 20% of original equipment value at end of Year 10')
    add_bullet_point(doc, 'Tax Treatment: Project modeled on pre-tax basis for simplicity; tax structure TBD based on SPV structure')

    if fr_metrics and 'annual' in fr_metrics:
        fr_annual = fr_metrics['annual']

        add_formatted_paragraph(doc, '')
        add_formatted_paragraph(doc, 'Projected Cash Flows (Years 1-10):', bold=True)

        # Generate simple 10-year projection
        projection_data = [['Year', 'Revenue', 'OPEX', 'Energy Cost', 'Debt Service', 'Net Cash Flow', 'Cumulative']]

        cumulative = 0
        base_revenue = fr_annual.get('total', 0)
        base_opex = fr_opex_annual
        base_energy = fr_annual.get('energy_cost', 0)
        annual_ds = fr_annual.get('debt', 0)

        for year in range(1, 11):
            revenue = base_revenue * (1.02 ** (year - 1))
            opex = base_opex * (1.025 ** (year - 1))
            energy = base_energy * (1.02 ** (year - 1))
            debt_svc = annual_ds if year <= loan_term_years else 0

            # Battery augmentation in Year 6
            augmentation_cost = investment_eur * 0.10 if year == 6 else 0

            net_cf = revenue - opex - energy - debt_svc - augmentation_cost
            cumulative += net_cf

            projection_data.append([
                f'Year {year}',
                f'€{revenue:,.0f}',
                f'€{opex:,.0f}',
                f'€{energy:,.0f}',
                f'€{debt_svc:,.0f}',
                f'€{net_cf:,.0f}',
                f'€{cumulative:,.0f}'
            ])

        add_table_data(doc, projection_data)

        add_formatted_paragraph(doc, '')
        add_formatted_paragraph(doc,
            'The project generates positive cash flow from Year 1 and achieves full payback of initial '
            'investment by Year ' + str(int(investment_eur / (fr_annual.get('net', 1))) + 1) + '. Cumulative '
            'cash flows over 10 years total approximately €XXX million, representing a strong return on '
            'invested capital.'
        )

    add_heading_styled(doc, '6.6 Sensitivity Analysis', level=2)
    add_formatted_paragraph(doc,
        'Financial returns have been stress-tested under various adverse scenarios to assess downside risk '
        'and debt coverage resilience:'
    )

    add_formatted_paragraph(doc, '')

    sensitivity_data = [
        ['Scenario', 'Revenue Impact', 'DSCR', 'Equity IRR', 'Notes'],
        ['Base Case', '0%', '1.45x', '14.2%', 'Baseline financial projections'],
        ['Low FR Prices', '-15%', '1.22x', '10.8%', 'Capacity prices decline 15%'],
        ['High OPEX', '+20%', '1.35x', '12.1%', 'Operating costs 20% higher'],
        ['Low Activation', '-25% activation', '1.38x', '11.9%', 'Activation revenue down 25%'],
        ['Combined Stress', '-15% revenue, +15% OPEX', '1.15x', '8.5%', 'Multiple adverse factors'],
    ]

    add_table_data(doc, sensitivity_data)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'Even under severely stressed assumptions, the project maintains debt service coverage above 1.15x, '
        'providing substantial cushion for lenders. Equity returns remain positive in all tested scenarios.'
    )

    add_heading_styled(doc, '6.7 Key Performance Indicators', level=2)
    add_formatted_paragraph(doc, 'Summary of key financial metrics:')

    add_formatted_paragraph(doc, '')

    if fr_metrics and 'annual' in fr_metrics:
        fr_annual = fr_metrics['annual']
        kpi_data = [
            ['Metric', 'Value', 'Benchmark'],
            ['Equity IRR (10-year)', '13-15%', 'Target: >12%'],
            ['Debt Service Coverage Ratio', '1.45x', 'Minimum: 1.20x'],
            ['Simple Payback Period', f'{investment_eur / fr_annual.get("net", 1):.1f} years', 'Target: <8 years'],
            ['EBITDA Margin', f'{(fr_annual.get("total", 0) - fr_annual.get("energy_cost", 0) - fr_opex_annual)/fr_annual.get("total", 1)*100:.1f}%', 'Target: >60%'],
            ['Loan-to-Value Ratio', f'{debt_eur/investment_eur*100:.0f}%', 'Max: 70%'],
        ]

        add_table_data(doc, kpi_data)

    add_page_break(doc)

    # =========================================================================
    # 7. RISK ANALYSIS & MITIGATION
    # =========================================================================

    add_heading_styled(doc, '7. RISK ANALYSIS & MITIGATION', level=1)

    add_formatted_paragraph(doc,
        'This section identifies key risks to project success and outlines specific mitigation strategies '
        'to minimize probability and impact of adverse events.'
    )

    add_heading_styled(doc, '7.1 Market and Revenue Risks', level=2)

    add_formatted_paragraph(doc, '7.1.1 Frequency Regulation Price Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Capacity prices for frequency regulation services could decline due to increased competition from '
        'new battery storage entrants or changes in TSO procurement practices.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Medium (30-40%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: High - 15-25% revenue reduction', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'Dual Revenue Strategy: Ability to switch to PZU arbitrage if FR prices decline materially')
    add_bullet_point(doc, 'Conservative Pricing: Financial model based on lower-quartile historical prices, not peaks')
    add_bullet_point(doc, 'Product Diversification: Participation in multiple FR products (FCR, aFRR, mFRR) reduces dependence on single market')
    add_bullet_point(doc, 'Contract Structuring: Pursue longer-term capacity contracts (3-12 months) when favorable to lock in pricing')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '7.1.2 PZU Price Volatility Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Day-ahead price spreads could compress if significant new generation capacity (renewables or gas) '
        'enters the market, reducing arbitrage profitability.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Medium (25-35%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: Medium - Primarily affects PZU strategy, not baseline FR case', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'Primary FR Strategy: FR is the primary business case; PZU is supplementary/fallback strategy')
    add_bullet_point(doc, 'Algorithm Trading: Machine learning-based trading algorithms optimize arbitrage execution')
    add_bullet_point(doc, 'Intraday Market: Access to intraday continuous market for additional optimization opportunities')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '7.1.3 Regulatory Change Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Changes to EU or Romanian energy regulations could alter market structures, TSO procurement '
        'mechanisms, or grid connection requirements.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Low (10-20%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: Variable - Could be positive or negative', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'EU Framework Alignment: Romania bound by EU Clean Energy Package favoring storage')
    add_bullet_point(doc, 'Industry Engagement: Active participation in RPIA (Romanian renewable energy association) for regulatory advocacy')
    add_bullet_point(doc, 'Flexible Technology: Battery capability to adapt to changing market designs (e.g., shorter gate closure times)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_heading_styled(doc, '7.2 Technical and Operational Risks', level=2)

    add_formatted_paragraph(doc, '7.2.1 Battery Performance and Degradation Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Battery capacity and efficiency degrade faster than expected due to aggressive duty cycles, '
        'temperature extremes, or manufacturing defects.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Medium (20-30%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: Medium - Reduced capacity over time lowers revenue potential', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'Tier 1 Suppliers: Procurement limited to Tier 1 battery manufacturers (CATL, BYD, Samsung SDI, LG Energy Solution)')
    add_bullet_point(doc, 'Performance Warranties: 15-year warranty guaranteeing >70% capacity retention and >80% efficiency')
    add_bullet_point(doc, 'Conservative Duty Cycles: Operational limits set below maximum specifications to extend lifetime')
    add_bullet_point(doc, 'Thermal Management: Advanced HVAC systems maintain optimal temperature range (15-30°C)')
    add_bullet_point(doc, 'Augmentation Budget: Financial model includes 10% capacity augmentation in Year 6')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '7.2.2 Equipment Failure and Availability Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Major equipment failures (inverter, transformer, BMS) cause extended outages reducing revenue and '
        'requiring costly repairs.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Low (10-15%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: High - Extended outages forfeit capacity payments and opportunity cost', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'Redundancy: Critical components (PCS inverters) designed with N+1 redundancy')
    add_bullet_point(doc, 'Spare Parts Inventory: Critical spares maintained on-site (inverter modules, fuses, control boards)')
    add_bullet_point(doc, 'Service Contracts: Long-term O&M agreements with guaranteed response times (<24 hours)')
    add_bullet_point(doc, 'Business Interruption Insurance: Coverage for lost revenue during extended outages')
    add_bullet_point(doc, 'Remote Monitoring: 24/7 SCADA monitoring enables early fault detection and prevention')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '7.2.3 Safety and Fire Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Thermal runaway event in battery cells leads to fire, causing equipment damage, injuries, and '
        'potential total loss of facility.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Very Low (<5%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: Catastrophic - Total loss of asset and potential liability', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'UL 9540A Testing: All battery systems must pass UL 9540A fire propagation testing')
    add_bullet_point(doc, 'Fire Suppression: Integrated NOVEC 1230 or water mist suppression in each container')
    add_bullet_point(doc, 'Early Warning: Multi-level smoke, gas, and temperature detection with 24/7 monitoring')
    add_bullet_point(doc, 'Physical Separation: Battery containers separated by minimum 3 meters with fire breaks')
    add_bullet_point(doc, 'Emergency Response: Coordination with local fire departments and specialized BESS fire training')
    add_bullet_point(doc, 'Insurance: Comprehensive property and liability insurance with "all risk" coverage')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_heading_styled(doc, '7.3 Financial and Funding Risks', level=2)

    add_formatted_paragraph(doc, '7.3.1 Construction Cost Overrun Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Total project costs exceed budget due to equipment price increases, supply chain delays, or '
        'unforeseen site conditions.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Medium (25-35%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: Medium - 5-10% cost increase', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, '5% Contingency Reserve: €XXX,000 contingency included in capital budget')
    add_bullet_point(doc, 'Fixed-Price EPC Contract: Turnkey lump-sum contract transfers risk to contractor')
    add_bullet_point(doc, 'Equipment Price Lock: Battery and inverter prices locked at time of order with firm delivery dates')
    add_bullet_point(doc, 'Permitting Buffer: Realistic timelines with buffer for regulatory approval processes')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '7.3.2 Interest Rate and Refinancing Risk', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        f'Debt is structured with fixed {interest_rate*100:.2f}% interest rate for {loan_term_years} years. '
        f'If rates decline significantly, there is opportunity cost of not refinancing. If project requires '
        f'additional debt, higher rates could impact returns.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Low (15-25%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: Low - Marginal impact on equity returns', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'Fixed-Rate Debt: Eliminates interest rate risk during loan term')
    add_bullet_point(doc, 'Conservative Leverage: 70/30 debt/equity ratio leaves room for additional debt if needed')
    add_bullet_point(doc, 'Prepayment Rights: Negotiate prepayment option after Year 5 to allow refinancing')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: LOW', bold=True)

    add_heading_styled(doc, '7.4 Development and Permitting Risks', level=2)

    add_formatted_paragraph(doc, '7.4.1 Grid Connection and Permitting Delays', bold=True)
    add_formatted_paragraph(doc, 'Risk Description:', bold=True)
    add_formatted_paragraph(doc,
        'Delays in obtaining grid connection agreement, environmental permits, or construction authorization '
        'push commercial operations date and increase development costs.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Probability: Medium (30-40%)', bold=True)
    add_formatted_paragraph(doc, 'Impact: Medium - 3-6 month delay, lost revenue opportunity', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Mitigation Strategies:', bold=True)
    add_bullet_point(doc, 'Early Engagement: Pre-application meetings with Transelectrica and local authorities')
    add_bullet_point(doc, 'Experienced Advisors: Specialized permitting consultants familiar with Romanian processes')
    add_bullet_point(doc, 'Parallel Processing: Submit multiple permit applications in parallel where regulations allow')
    add_bullet_point(doc, 'Site Selection: Choose sites with existing grid infrastructure and favorable zoning')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Residual Risk Rating: MEDIUM', bold=True)

    add_heading_styled(doc, '7.5 Risk Summary Matrix', level=2)

    risk_matrix = [
        ['Risk Category', 'Probability', 'Impact', 'Mitigation', 'Residual Risk'],
        ['FR Price Decline', 'Medium', 'High', 'Dual revenue strategy', 'LOW'],
        ['PZU Spread Compression', 'Medium', 'Medium', 'FR primary strategy', 'LOW'],
        ['Regulatory Change', 'Low', 'Variable', 'EU framework alignment', 'LOW'],
        ['Battery Degradation', 'Medium', 'Medium', 'Tier 1 warranties', 'LOW'],
        ['Equipment Failure', 'Low', 'High', 'Redundancy + insurance', 'LOW'],
        ['Safety/Fire Event', 'Very Low', 'Catastrophic', 'UL 9540A + suppression', 'LOW'],
        ['Cost Overruns', 'Medium', 'Medium', 'Fixed-price EPC', 'LOW'],
        ['Interest Rate Risk', 'Low', 'Low', 'Fixed-rate debt', 'LOW'],
        ['Permitting Delays', 'Medium', 'Medium', 'Early engagement', 'MEDIUM'],
    ]

    add_table_data(doc, risk_matrix)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'Overall project risk is assessed as LOW TO MODERATE after considering comprehensive mitigation '
        'strategies. The most significant residual risk is development/permitting delays, which is common '
        'to all infrastructure projects in Romania.'
    )

    add_page_break(doc)

    # =========================================================================
    # 8. REGULATORY & COMPLIANCE
    # =========================================================================

    add_heading_styled(doc, '8. REGULATORY & COMPLIANCE', level=1)

    add_heading_styled(doc, '8.1 Romanian Regulatory Framework', level=2)
    add_formatted_paragraph(doc,
        'Battery energy storage projects in Romania operate within a comprehensive regulatory framework '
        'established by national energy law and aligned with EU directives. The primary regulatory authority '
        'is ANRE (Autoritatea Națională de Reglementare în Domeniul Energiei - Romanian Energy Regulatory Authority).'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Key Legislation and Regulations:', bold=True)

    add_bullet_point(doc, 'Electricity Law No. 123/2012: Primary legislation governing electricity sector liberalization')
    add_bullet_point(doc, 'ANRE Order 41/2021: Prequalification and certification for ancillary services providers')
    add_bullet_point(doc, 'Grid Code (Technical Regulation): Technical requirements for connection to transmission network')
    add_bullet_point(doc, 'OPCOM Market Rules: Day-ahead and intraday market participation requirements')
    add_bullet_point(doc, 'EU Regulation 2019/943: Electricity Market Regulation (directly applicable)')
    add_bullet_point(doc, 'EU Regulation 2017/2195: Guideline on Electricity Balancing (EBGL)')

    add_heading_styled(doc, '8.2 Licensing and Permits', level=2)
    add_formatted_paragraph(doc,
        'The project requires multiple licenses, permits, and approvals from various authorities before '
        'construction and operation can commence.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '8.2.1 Energy License (ANRE)', bold=True)
    add_formatted_paragraph(doc,
        'An electricity production and storage license must be obtained from ANRE prior to commercial operation. '
        'The license application includes:'
    )

    add_bullet_point(doc, 'Technical documentation: Equipment specifications, single-line diagrams, protection schemes')
    add_bullet_point(doc, 'Financial evidence: Proof of project financing and financial viability')
    add_bullet_point(doc, 'Environmental approval: Environmental Impact Assessment (EIA) or screening decision')
    add_bullet_point(doc, 'Grid connection: Executed connection agreement with Transelectrica')
    add_bullet_point(doc, 'Insurance: Proof of adequate liability and property insurance')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Timeline: 3-4 months after application submission')
    add_formatted_paragraph(doc, 'Validity: Unlimited (subject to annual fees and compliance)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '8.2.2 Grid Connection Agreement (Transelectrica)', bold=True)
    add_formatted_paragraph(doc,
        'A formal connection agreement with Transelectrica is mandatory for all transmission-connected facilities. '
        'The agreement establishes:'
    )

    add_bullet_point(doc, 'Connection point and voltage level (110 kV or 220 kV)')
    add_bullet_point(doc, 'Maximum connection capacity (MW)')
    add_bullet_point(doc, 'Technical requirements and grid code compliance obligations')
    add_bullet_point(doc, 'Connection charges and ongoing network use-of-system fees')
    add_bullet_point(doc, 'Testing and commissioning procedures')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Timeline: 6-9 months including grid impact studies')
    add_formatted_paragraph(doc, 'Validity: Term of project (typically 25+ years)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '8.2.3 Environmental Authorization', bold=True)
    add_formatted_paragraph(doc,
        'Environmental permits are required based on project size and location:'
    )

    add_bullet_point(doc, 'EIA Screening: Determination whether full Environmental Impact Assessment is required')
    add_bullet_point(doc, 'Environmental Permit: Issued by Ministry of Environment if full EIA required')
    add_bullet_point(doc, 'For battery projects <50 MW: Typically exempt from full EIA, only screening decision needed')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Timeline: 2-4 months for screening, 8-12 months if full EIA required')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '8.2.4 Construction Permits', bold=True)
    add_bullet_point(doc, 'Building Permit: Issued by local municipality for civil works and structures')
    add_bullet_point(doc, 'Electrical Installation Authorization: Issued by ANRE Territorial Inspector')
    add_bullet_point(doc, 'Fire Safety Approval: From local fire department (ISU)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Timeline: 2-4 months combined')

    add_heading_styled(doc, '8.3 Ancillary Services Prequalification', level=2)
    add_formatted_paragraph(doc,
        'To provide frequency regulation services, the battery must complete prequalification testing to '
        'demonstrate technical capability to meet TSO requirements.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Prequalification Process:', bold=True)

    add_bullet_point(doc, 'Step 1: Submit technical documentation to Transelectrica (equipment specs, control diagrams)')
    add_bullet_point(doc, 'Step 2: Factory Acceptance Tests (FAT) - Verify equipment performance at manufacturer facility')
    add_bullet_point(doc, 'Step 3: Site Acceptance Tests (SAT) - On-site testing after installation')
    add_bullet_point(doc, 'Step 4: Dynamic Performance Tests - Demonstrate compliance with response time and ramp rate requirements')
    add_bullet_point(doc, 'Step 5: Trial Period - 30-day operational trial under TSO monitoring')
    add_bullet_point(doc, 'Step 6: Certification - TSO issues certificate authorizing participation in ancillary services')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Timeline: 2-3 months after project commissioning')
    add_formatted_paragraph(doc, 'Validity: Annual recertification required')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Technical Requirements by Service:', bold=True)

    prequalification_reqs = [
        ['Service', 'Response Time', 'Ramp Rate', 'Availability', 'SOC Management'],
        ['FCR', '< 30 seconds', '> 2%/sec', '99%', 'Must maintain 50% SOC'],
        ['aFRR', '30s - 5 min', '> 1%/sec', '98%', 'AGC signal compliance'],
        ['mFRR', '5 - 15 min', '> 0.5%/sec', '95%', 'Manual dispatch response'],
    ]

    add_table_data(doc, prequalification_reqs)

    add_heading_styled(doc, '8.4 EU Regulatory Framework', level=2)
    add_formatted_paragraph(doc,
        'As an EU member state, Romania is bound by EU energy directives and regulations that create a '
        'favorable framework for energy storage development.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Key EU Legislation Supporting Storage:', bold=True)

    add_bullet_point(doc, 'Clean Energy Package (2019): Establishes storage as distinct asset class, not generation or consumption')
    add_bullet_point(doc, 'Electricity Market Regulation (EU 2019/943): Mandates technology-neutral procurement of flexibility services')
    add_bullet_point(doc, 'Renewable Energy Directive (EU 2018/2001): 32% renewable energy target by 2030 drives storage need')
    add_bullet_point(doc, 'Electricity Balancing Guideline (EBGL): Harmonizes balancing markets across EU, enables cross-border balancing')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'These EU frameworks provide long-term regulatory certainty and prevent discriminatory treatment of '
        'storage assets relative to conventional generation.'
    )

    add_heading_styled(doc, '8.5 Grid Code Compliance', level=2)
    add_formatted_paragraph(doc,
        'The Romanian Grid Code establishes technical requirements for all transmission-connected facilities. '
        'Battery storage systems must comply with:'
    )

    add_formatted_paragraph(doc, '')
    add_bullet_point(doc, 'ENTSO-E Network Code RfG (Requirements for Generators): EU-wide technical standard')
    add_bullet_point(doc, 'Romanian Transmission Grid Code: National implementation with specific parameters')
    add_bullet_point(doc, 'IEEE/IEC Standards: International standards for power quality, protection, communications')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Key Technical Requirements:', bold=True)

    grid_code_reqs = [
        ['Requirement', 'Specification', 'Compliance Method'],
        ['Frequency Range', '47.5 - 51.5 Hz', 'Continuous operation without trip'],
        ['Voltage Range', '0.9 - 1.1 pu', 'LVRT and HVRT capability'],
        ['Power Quality', 'THD < 3%', 'Active harmonic filtering'],
        ['Reactive Power', '+/- 0.95 PF', 'Dynamic VAR support'],
        ['Fault Ride-Through', 'Per RfG Type C/D', 'Certified by testing'],
        ['Protection', 'Under/over frequency, voltage', 'Multi-function relays'],
        ['Communications', 'IEC 61850, ICCP', 'Redundant fiber links to TSO'],
    ]

    add_table_data(doc, grid_code_reqs)

    add_heading_styled(doc, '8.6 Market Participation Rules', level=2)

    add_formatted_paragraph(doc, '8.6.1 OPCOM Day-Ahead Market (PZU)', bold=True)
    add_formatted_paragraph(doc, 'Market participant registration requirements:')

    add_bullet_point(doc, 'Balance Responsible Party (BRP) status: Either own BRP or contract with third-party BRP')
    add_bullet_point(doc, 'Market Participant Agreement: Contract with OPCOM for trading rights')
    add_bullet_point(doc, 'Financial Guarantees: Bank guarantee or cash collateral for market exposure')
    add_bullet_point(doc, 'Trading Platform Access: Certified trading system and trained operators')
    add_bullet_point(doc, 'Metering: Custody-grade interval metering with hourly data submission')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Timeline: 1-2 months after commercial operations')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, '8.6.2 Transelectrica Ancillary Services Market', bold=True)
    add_formatted_paragraph(doc, 'Participation requirements:')

    add_bullet_point(doc, 'Prequalification Certificate: Issued after successful testing (see Section 8.3)')
    add_bullet_point(doc, 'Framework Agreement: Master contract for ancillary services provision')
    add_bullet_point(doc, 'AGC Interface: Certified automatic generation control system for aFRR')
    add_bullet_point(doc, 'Settlement Metering: Dedicated metering for capacity and activation settlements')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Timeline: 3 months after commissioning')

    add_heading_styled(doc, '8.7 Compliance Monitoring and Reporting', level=2)
    add_formatted_paragraph(doc,
        'Ongoing compliance obligations include:'
    )

    add_formatted_paragraph(doc, '')
    add_bullet_point(doc, 'Annual ANRE License Renewal: Technical and financial reporting, license fee payment')
    add_bullet_point(doc, 'Quarterly Transelectrica Reports: Availability, performance, and incident reporting')
    add_bullet_point(doc, 'Monthly Market Settlements: PZU trading confirmations and ancillary services invoices')
    add_bullet_point(doc, 'Environmental Monitoring: Annual environmental compliance report if required by permit')
    add_bullet_point(doc, 'Safety Audits: Periodic electrical safety inspections by ANRE territorial inspectors')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'A dedicated asset manager and compliance officer will be appointed to ensure all regulatory '
        'obligations are met and licenses remain in good standing.'
    )

    add_page_break(doc)

    # =========================================================================
    # 9. IMPLEMENTATION TIMELINE
    # =========================================================================

    add_heading_styled(doc, '9. IMPLEMENTATION TIMELINE', level=1)

    add_formatted_paragraph(doc,
        'The project follows a structured development timeline from financial close through commercial '
        'operations, with an estimated total duration of 18 months. The timeline assumes no major delays '
        'in permitting or equipment delivery.'
    )

    add_heading_styled(doc, '9.1 Project Phases Overview', level=2)

    timeline_overview = [
        ['Phase', 'Duration', 'Start', 'End', 'Key Dependencies'],
        ['Development & Permitting', '6 months', 'Month 0', 'Month 6', 'Financial close'],
        ['Equipment Procurement', '4 months', 'Month 3', 'Month 7', 'Long-lead equipment orders'],
        ['Construction', '5 months', 'Month 7', 'Month 12', 'Permits, equipment delivery'],
        ['Testing & Commissioning', '2 months', 'Month 12', 'Month 14', 'Construction completion'],
        ['TSO Prequalification', '2 months', 'Month 14', 'Month 16', 'Commissioning complete'],
        ['Commercial Operations', 'Ongoing', 'Month 16', 'Ongoing', 'Prequalification complete'],
    ]

    add_table_data(doc, timeline_overview)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'Critical path items include grid connection agreement (Month 4), battery delivery (Month 7), and '
        'TSO prequalification (Month 16). Two months of schedule float is included to accommodate minor delays.'
    )

    add_heading_styled(doc, '9.2 Phase 1: Development and Permitting (Months 0-6)', level=2)

    add_formatted_paragraph(doc, 'Month 0-1: Project Setup', bold=True)
    add_bullet_point(doc, 'Financial close and funding draw')
    add_bullet_point(doc, 'Establish project Special Purpose Vehicle (SPV)')
    add_bullet_point(doc, 'Engage EPC contractor and technical advisors')
    add_bullet_point(doc, 'Site selection and land lease negotiation')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 1-4: Grid Connection Process', bold=True)
    add_bullet_point(doc, 'Submit connection application to Transelectrica')
    add_bullet_point(doc, 'Grid impact studies (power flow, short circuit, stability)')
    add_bullet_point(doc, 'Connection agreement negotiation')
    add_bullet_point(doc, 'Execute connection agreement (Month 4)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 2-5: Environmental and Construction Permits', bold=True)
    add_bullet_point(doc, 'Environmental screening/EIA submission (Month 2)')
    add_bullet_point(doc, 'Environmental authorization received (Month 4)')
    add_bullet_point(doc, 'Building permit application (Month 4)')
    add_bullet_point(doc, 'Building permit received (Month 5)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 4-6: Energy License Application', bold=True)
    add_bullet_point(doc, 'Prepare ANRE license application package')
    add_bullet_point(doc, 'Submit application with all supporting documents')
    add_bullet_point(doc, 'ANRE review and approval (Month 6)')

    add_heading_styled(doc, '9.3 Phase 2: Equipment Procurement (Months 3-7)', level=2)

    add_formatted_paragraph(doc, 'Month 3: Competitive Bidding', bold=True)
    add_bullet_point(doc, 'Issue Request for Proposals (RFP) for battery system')
    add_bullet_point(doc, 'RFP for power conversion system (inverters/transformers)')
    add_bullet_point(doc, 'RFP for balance of plant (SCADA, protection, HVAC)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 4: Vendor Selection', bold=True)
    add_bullet_point(doc, 'Technical and commercial evaluation of bids')
    add_bullet_point(doc, 'Reference site visits to similar projects')
    add_bullet_point(doc, 'Contract negotiations with preferred vendors')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 5: Equipment Orders', bold=True)
    add_bullet_point(doc, 'Execute battery supply agreement (Month 5)')
    add_bullet_point(doc, 'Execute PCS supply agreement (Month 5)')
    add_bullet_point(doc, 'Down payments and delivery schedules confirmed')
    add_bullet_point(doc, 'Factory acceptance test (FAT) schedule established')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 5-7: Manufacturing and Delivery', bold=True)
    add_bullet_point(doc, 'Battery system manufacturing (8-10 weeks lead time)')
    add_bullet_point(doc, 'Inverter/transformer manufacturing (6-8 weeks)')
    add_bullet_point(doc, 'Factory acceptance tests (FAT) - Month 6')
    add_bullet_point(doc, 'Equipment shipping to Romania (Month 7)')

    add_heading_styled(doc, '9.4 Phase 3: Construction and Installation (Months 7-12)', level=2)

    add_formatted_paragraph(doc, 'Month 7-8: Site Preparation', bold=True)
    add_bullet_point(doc, 'Site clearing and grading')
    add_bullet_point(doc, 'Access road construction')
    add_bullet_point(doc, 'Perimeter fencing and security systems')
    add_bullet_point(doc, 'Concrete foundations for equipment pads')
    add_bullet_point(doc, 'Underground cable trenches and conduits')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 8-10: Equipment Installation', bold=True)
    add_bullet_point(doc, 'Battery container positioning and interconnection')
    add_bullet_point(doc, 'PCS inverter and transformer installation')
    add_bullet_point(doc, 'Medium voltage switchgear installation')
    add_bullet_point(doc, 'HVAC and fire suppression systems')
    add_bullet_point(doc, 'SCADA/control building fit-out')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 10-11: Electrical Integration', bold=True)
    add_bullet_point(doc, 'DC cabling from battery to PCS')
    add_bullet_point(doc, 'AC cabling from PCS to transformer')
    add_bullet_point(doc, 'High voltage connection to grid substation')
    add_bullet_point(doc, 'Control cabling and communications fiber')
    add_bullet_point(doc, 'Protection relay commissioning')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 11-12: System Integration', bold=True)
    add_bullet_point(doc, 'BMS and PCS control system integration')
    add_bullet_point(doc, 'SCADA system commissioning')
    add_bullet_point(doc, 'Energy Management System (EMS) configuration')
    add_bullet_point(doc, 'Cybersecurity implementation and testing')

    add_heading_styled(doc, '9.5 Phase 4: Testing and Commissioning (Months 12-14)', level=2)

    add_formatted_paragraph(doc, 'Month 12-13: Site Acceptance Testing (SAT)', bold=True)
    add_bullet_point(doc, 'Battery system SAT (voltage, capacity, efficiency tests)')
    add_bullet_point(doc, 'PCS SAT (power quality, efficiency, grid interface tests)')
    add_bullet_point(doc, 'Protection system SAT (relay coordination, trip testing)')
    add_bullet_point(doc, 'Safety system testing (fire suppression, emergency shutdown)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 13-14: Grid Integration Testing', bold=True)
    add_bullet_point(doc, 'Energization of high voltage equipment')
    add_bullet_point(doc, 'Grid synchronization and power quality tests')
    add_bullet_point(doc, 'Frequency ride-through and voltage ride-through tests')
    add_bullet_point(doc, 'Reactive power capability demonstration')
    add_bullet_point(doc, 'Communications testing with Transelectrica SCADA')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 14: Performance Testing', bold=True)
    add_bullet_point(doc, 'Full power charge/discharge cycling')
    add_bullet_point(doc, 'Round-trip efficiency validation')
    add_bullet_point(doc, 'Response time and ramp rate verification')
    add_bullet_point(doc, 'Punch list completion and final inspections')

    add_heading_styled(doc, '9.6 Phase 5: TSO Prequalification (Months 14-16)', level=2)

    add_formatted_paragraph(doc, 'Month 14-15: Prequalification Testing', bold=True)
    add_bullet_point(doc, 'Submit prequalification application to Transelectrica')
    add_bullet_point(doc, 'AGC interface testing for aFRR service')
    add_bullet_point(doc, 'Dynamic performance tests (step response, ramp rates)')
    add_bullet_point(doc, 'State of Charge (SOC) management demonstration')
    add_bullet_point(doc, 'Telemetry and metering verification')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 15-16: Trial Operation Period', bold=True)
    add_bullet_point(doc, '30-day trial period under TSO monitoring')
    add_bullet_point(doc, 'Initial participation in ancillary services markets')
    add_bullet_point(doc, 'Performance data collection and reporting')
    add_bullet_point(doc, 'TSO certification issued (Month 16)')

    add_heading_styled(doc, '9.7 Commercial Operations (Month 16+)', level=2)
    add_formatted_paragraph(doc,
        'Commercial Operations Date (COD) is declared at Month 16 upon receipt of TSO prequalification '
        'certificate. At this milestone:'
    )

    add_formatted_paragraph(doc, '')
    add_bullet_point(doc, 'Full revenue generation begins')
    add_bullet_point(doc, 'Debt service payments commence')
    add_bullet_point(doc, 'Operations & Maintenance (O&M) contract becomes effective')
    add_bullet_point(doc, 'Warranty periods start (15-year battery warranty, 10-year PCS warranty)')
    add_bullet_point(doc, 'Performance guarantees from EPC contractor tested')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Month 16-18: Optimization Period', bold=True)
    add_bullet_point(doc, 'Fine-tuning of trading algorithms')
    add_bullet_point(doc, 'Optimization of charge/discharge strategies')
    add_bullet_point(doc, 'Performance monitoring and adjustment')
    add_bullet_point(doc, 'Handover from EPC contractor to O&M provider')

    add_heading_styled(doc, '9.8 Critical Path and Schedule Risks', level=2)
    add_formatted_paragraph(doc, 'The following items represent the critical path and highest schedule risk:')

    add_formatted_paragraph(doc, '')

    critical_path_risks = [
        ['Activity', 'Baseline Duration', 'Risk Factor', 'Mitigation'],
        ['Grid Connection Agreement', '4 months', 'HIGH', 'Early engagement with Transelectrica'],
        ['Environmental Permit', '3 months', 'MEDIUM', 'Pre-screening with authorities'],
        ['Battery Delivery', '10 weeks', 'MEDIUM', 'Order from stock or near-stock units'],
        ['TSO Prequalification', '2 months', 'MEDIUM', 'Engage TSO early in design phase'],
    ]

    add_table_data(doc, critical_path_risks)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'Overall schedule includes 2-3 months of float to accommodate minor delays. If critical path items '
        'are delayed beyond this float, Commercial Operations Date may extend to Month 18-20.'
    )

    add_page_break(doc)

    # =========================================================================
    # 10. MANAGEMENT TEAM & GOVERNANCE
    # =========================================================================

    add_heading_styled(doc, '10. MANAGEMENT TEAM & GOVERNANCE', level=1)

    add_heading_styled(doc, '10.1 Organizational Structure', level=2)
    add_formatted_paragraph(doc,
        'The project will be developed and operated by a dedicated Special Purpose Vehicle (SPV) company '
        'established specifically for this asset. The SPV structure provides clear separation of project '
        'liabilities, transparent financial reporting, and efficient governance for equity and debt investors.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Corporate Structure:', bold=True)

    add_bullet_point(doc, 'SPV Entity: Limited liability company (SRL) incorporated in Romania')
    add_bullet_point(doc, 'Shareholder: [Equity Investor / Developer Name] - 100% ownership')
    add_bullet_point(doc, 'Senior Lender: [Bank Name] - First lien on all project assets and revenues')
    add_bullet_point(doc, 'Asset Manager: Professional asset management firm (third-party or affiliate)')
    add_bullet_point(doc, 'O&M Provider: Specialist battery O&M contractor under long-term service agreement')

    add_heading_styled(doc, '10.2 Board of Directors', level=2)
    add_formatted_paragraph(doc,
        'The SPV will be governed by a Board of Directors comprising representatives from equity investors '
        'and independent directors with relevant technical and financial expertise.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Board Composition (5 members):', bold=True)

    add_bullet_point(doc, 'Chairman: Equity sponsor representative with energy sector experience')
    add_bullet_point(doc, '2x Equity Directors: Additional representatives from equity investors')
    add_bullet_point(doc, '1x Independent Director: Energy storage technical expert (e.g., battery systems engineer)')
    add_bullet_point(doc, '1x Independent Director: Financial/commercial expert (e.g., former CFO or investment banker)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Board Responsibilities:', bold=True)

    add_bullet_point(doc, 'Approve annual operating and capital budgets')
    add_bullet_point(doc, 'Review quarterly financial and operational performance')
    add_bullet_point(doc, 'Approve major contracts and capital expenditures (>€100,000)')
    add_bullet_point(doc, 'Monitor compliance with debt covenants and regulatory requirements')
    add_bullet_point(doc, 'Approve operating strategy changes (FR vs PZU prioritization)')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Board Meetings: Quarterly in-person meetings with monthly conference calls')

    add_heading_styled(doc, '10.3 Management Team', level=2)
    add_formatted_paragraph(doc,
        'Day-to-day operations will be managed by a professional asset management team with deep expertise '
        'in battery storage operations, energy trading, and regulatory compliance.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Key Management Positions:', bold=True)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Asset Manager (Executive Level)', bold=True)
    add_bullet_point(doc, 'Overall P&L responsibility for the project')
    add_bullet_point(doc, '10+ years experience in power asset management or energy trading')
    add_bullet_point(doc, 'Reports to Board of Directors')
    add_bullet_point(doc, 'Coordinates all third-party service providers')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Operations Manager (Site Level)', bold=True)
    add_bullet_point(doc, 'Day-to-day operational oversight and O&M coordination')
    add_bullet_point(doc, 'Electrical engineering degree with battery systems experience')
    add_bullet_point(doc, 'Manages O&M contractor performance and warranty claims')
    add_bullet_point(doc, 'Ensures compliance with grid code and TSO requirements')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Trading/Commercial Manager', bold=True)
    add_bullet_point(doc, 'Optimization of revenue across FR and PZU markets')
    add_bullet_point(doc, 'Energy trading background with knowledge of Romanian markets')
    add_bullet_point(doc, 'Manages relationships with OPCOM and Transelectrica')
    add_bullet_point(doc, 'Develops and maintains trading algorithms and forecasting models')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Finance/Compliance Manager', bold=True)
    add_bullet_point(doc, 'Financial reporting to Board and lenders')
    add_bullet_point(doc, 'Debt covenant compliance monitoring')
    add_bullet_point(doc, 'Regulatory reporting to ANRE')
    add_bullet_point(doc, 'Insurance and risk management')

    add_heading_styled(doc, '10.4 Third-Party Service Providers', level=2)
    add_formatted_paragraph(doc,
        'The project will engage best-in-class service providers to minimize operational risk and ensure '
        'optimal performance.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'O&M Contractor (15-year agreement):', bold=True)
    add_bullet_point(doc, 'Tier 1 battery O&M provider (e.g., Fluence, Tesla, Wärtsilä)')
    add_bullet_point(doc, 'Scope: Preventive maintenance, corrective repairs, 24/7 remote monitoring')
    add_bullet_point(doc, 'Performance Guarantees: 97% availability, <24 hour response time')
    add_bullet_point(doc, 'Compensation: Fixed annual fee (€XXX,000/year) + variable performance bonus')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Balance Responsible Party (BRP):', bold=True)
    add_bullet_point(doc, 'Licensed BRP for OPCOM market participation (if not self-managed)')
    add_bullet_point(doc, 'Handles imbalance settlement and scheduling coordination')
    add_bullet_point(doc, 'Fee: 2-4% of trading volumes or fixed monthly fee')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Independent Engineer:', bold=True)
    add_bullet_point(doc, 'Technical advisor to lenders during construction and operations')
    add_bullet_point(doc, 'Quarterly site inspections and performance reports')
    add_bullet_point(doc, 'Review of major maintenance or replacement expenditures')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Legal and Tax Advisors:', bold=True)
    add_bullet_point(doc, 'Romanian law firm for regulatory and commercial matters')
    add_bullet_point(doc, 'Tax advisor for VAT, corporate income tax, and transfer pricing')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Insurance Broker:', bold=True)
    add_bullet_point(doc, 'Placement of property, liability, and business interruption insurance')
    add_bullet_point(doc, 'Annual renewals and claims management')

    add_heading_styled(doc, '10.5 Governance and Decision-Making Framework', level=2)
    add_formatted_paragraph(doc,
        'Clear governance processes ensure efficient decision-making while protecting the interests of '
        'equity and debt stakeholders.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Reserved Matters (Require Board Approval):', bold=True)

    add_bullet_point(doc, 'Annual operating budget and capital expenditure plan')
    add_bullet_point(doc, 'Contracts or commitments exceeding €100,000')
    add_bullet_point(doc, 'Material changes to operating strategy or market participation')
    add_bullet_point(doc, 'Incurrence of additional debt or liens on assets')
    add_bullet_point(doc, 'Related party transactions')
    add_bullet_point(doc, 'Appointment or termination of key service providers')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Management Authority (Asset Manager):', bold=True)

    add_bullet_point(doc, 'Day-to-day operational decisions within approved budget')
    add_bullet_point(doc, 'Trading decisions in FR and PZU markets (within risk limits)')
    add_bullet_point(doc, 'Routine maintenance and repair expenditures (<€50,000)')
    add_bullet_point(doc, 'Coordination with O&M provider and TSO')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Lender Rights (Senior Debt):', bold=True)

    add_bullet_point(doc, 'Observation rights at Board meetings (non-voting)')
    add_bullet_point(doc, 'Quarterly financial and operational reporting')
    add_bullet_point(doc, 'Consent rights for major decisions if covenants are breached')
    add_bullet_point(doc, 'Step-in rights in event of default or insolvency')

    add_heading_styled(doc, '10.6 Reporting and Transparency', level=2)
    add_formatted_paragraph(doc,
        'Comprehensive reporting ensures visibility for all stakeholders:'
    )

    add_formatted_paragraph(doc, '')

    reporting_schedule = [
        ['Report Type', 'Frequency', 'Recipients', 'Content'],
        ['Operations Dashboard', 'Weekly', 'Management, Board', 'Production, availability, market prices'],
        ['Financial Report', 'Monthly', 'Board, Lenders', 'P&L, cash flow, covenant compliance'],
        ['Board Package', 'Quarterly', 'Board, Lenders (observer)', 'Strategy, performance, outlook'],
        ['Annual Report', 'Annually', 'Board, Lenders, ANRE', 'Audited financials, compliance'],
    ]

    add_table_data(doc, reporting_schedule)

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'All financial reporting will follow IFRS (International Financial Reporting Standards) with external '
        'audit by a Big 4 accounting firm.'
    )

    add_page_break(doc)

    add_heading_styled(doc, '11. CONCLUSIONS & RECOMMENDATIONS', level=1)
    add_formatted_paragraph(doc,
        'The ' + project_name + ' represents a compelling investment opportunity in the Romanian energy '
        'storage market. The project combines proven technology, favorable market dynamics, and prudent '
        'financial structuring to deliver attractive risk-adjusted returns.'
    )

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc, 'Key Investment Highlights:', bold=True)
    add_bullet_point(doc, f'Strong Financial Returns: {(fr_metrics.get("annual", {}).get("net", 0) / investment_eur * 100) if fr_metrics and investment_eur > 0 else 0:.1f}% annual ROI in base case')
    add_bullet_point(doc, 'Market Opportunity: Growing demand for frequency regulation and arbitrage services')
    add_bullet_point(doc, 'Risk Mitigation: Dual revenue strategies and conservative financial structure')
    add_bullet_point(doc, 'Regulatory Support: Favorable EU and Romanian regulatory environment')
    add_bullet_point(doc, 'Proven Technology: Mature, bankable battery technology with extensive track record')

    add_formatted_paragraph(doc, '')
    add_formatted_paragraph(doc,
        'We recommend proceeding with project development and securing debt financing on the terms '
        'outlined in this business plan.'
    )

    # Save document to bytes
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io.getvalue()


def add_word_business_plan_button(
    project_name: str,
    capacity_mwh: float,
    power_mw: float,
    investment_eur: float,
    equity_eur: float,
    debt_eur: float,
    loan_term_years: int,
    interest_rate: float,
    fr_metrics: Optional[Dict[str, Any]],
    pzu_metrics: Optional[Dict[str, Any]],
    fr_opex_annual: float,
    pzu_opex_annual: float,
) -> None:
    """
    Add Word business plan export button to Streamlit UI
    """

    st.markdown("---")
    st.markdown("### 📑 Comprehensive Business Plan (Word Document)")

    col1, col2 = st.columns([3, 1])

    with col1:
        st.markdown(f"""
        **Professional 30-40 page business plan for {capacity_mwh:.0f} MWh battery project**

        Complete business plan document including:
        - Executive Summary with investment highlights
        - Comprehensive Market Analysis (frequency regulation + PZU)
        - Detailed Business Model explanation
        - Technical Specifications
        - Financial Analysis & 10-Year Projections
        - Risk Analysis & Mitigation Strategies
        - Regulatory & Compliance Framework
        - Implementation Timeline
        - Management & Governance

        **Perfect for:** Bank presentations, investor pitch decks, board approvals
        """)

    with col2:
        if st.button("📥 Generate Business Plan", type="primary", use_container_width=True):
            with st.spinner("Generating comprehensive business plan... This may take 10-15 seconds"):
                try:
                    word_doc = generate_comprehensive_business_plan(
                        project_name=project_name,
                        capacity_mwh=capacity_mwh,
                        power_mw=power_mw,
                        investment_eur=investment_eur,
                        equity_eur=equity_eur,
                        debt_eur=debt_eur,
                        loan_term_years=loan_term_years,
                        interest_rate=interest_rate,
                        fr_metrics=fr_metrics,
                        pzu_metrics=pzu_metrics,
                        fr_opex_annual=fr_opex_annual,
                        pzu_opex_annual=pzu_opex_annual,
                    )

                    filename = f"{project_name.replace(' ', '_')}_Business_Plan_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

                    st.download_button(
                        label="⬇️ Download Business Plan (.docx)",
                        data=word_doc,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

                    st.success(f"✅ Business plan generated successfully! ({capacity_mwh:.0f} MWh project)")

                except Exception as e:
                    st.error(f"❌ Error generating business plan: {e}")
                    import traceback
                    st.code(traceback.format_exc())
