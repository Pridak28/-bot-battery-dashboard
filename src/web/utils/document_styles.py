"""
Professional Document Styling for Excel and Word Exports

Provides consistent, investment banking-grade visual styling for:
- Excel financial models (JP Morgan, Goldman Sachs style)
- Word business plans (McKinsey, BCG style)

Color scheme follows professional financial services standards.
"""

from typing import Dict, Any
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# ============================================================================
# COLOR PALETTE - Investment Banking Standard
# ============================================================================

class DocumentColors:
    """Professional color palette for financial documents"""

    # Primary colors (JP Morgan / Goldman Sachs style)
    NAVY_BLUE = "1F4E78"        # Dark navy for titles
    MEDIUM_BLUE = "4472C4"      # Medium blue for headers
    LIGHT_BLUE = "D9E1F2"       # Light blue for totals/highlights
    ACCENT_BLUE = "2E75B6"      # Accent blue for charts

    # Neutral colors
    WHITE = "FFFFFF"
    LIGHT_GRAY = "F2F2F2"       # Alternating rows
    MEDIUM_GRAY = "A6A6A6"      # Borders/dividers
    DARK_GRAY = "404040"        # Secondary text
    BLACK = "000000"

    # Status colors
    SUCCESS_GREEN = "548235"    # Positive metrics
    WARNING_ORANGE = "C65911"   # Caution items
    ERROR_RED = "C00000"        # Negative/risk items

    # RGB equivalents for Word documents
    NAVY_BLUE_RGB = RGBColor(31, 78, 120)
    MEDIUM_BLUE_RGB = RGBColor(68, 114, 196)
    LIGHT_BLUE_RGB = RGBColor(217, 225, 242)
    DARK_GRAY_RGB = RGBColor(64, 64, 64)
    SUCCESS_GREEN_RGB = RGBColor(84, 130, 53)
    WHITE_RGB = RGBColor(255, 255, 255)


# ============================================================================
# FONT SETTINGS
# ============================================================================

class DocumentFonts:
    """Professional typography for financial documents"""

    # Excel fonts
    EXCEL_TITLE_FONT = "Calibri"
    EXCEL_HEADER_FONT = "Calibri"
    EXCEL_BODY_FONT = "Calibri"
    EXCEL_NUMBER_FONT = "Calibri"

    # Word fonts
    WORD_TITLE_FONT = "Calibri"
    WORD_HEADING_FONT = "Calibri"
    WORD_BODY_FONT = "Calibri"
    WORD_TABLE_FONT = "Calibri"

    # Font sizes
    EXCEL_TITLE_SIZE = 14
    EXCEL_HEADER_SIZE = 11
    EXCEL_BODY_SIZE = 10

    WORD_TITLE_SIZE = Pt(28)
    WORD_HEADING1_SIZE = Pt(18)
    WORD_HEADING2_SIZE = Pt(14)
    WORD_HEADING3_SIZE = Pt(12)
    WORD_BODY_SIZE = Pt(11)
    WORD_TABLE_SIZE = Pt(10)


# ============================================================================
# EXCEL STYLING FUNCTIONS
# ============================================================================

def get_excel_title_style() -> Dict[str, Any]:
    """Get Excel title row style (navy blue background, white text)"""
    return {
        'font': Font(
            bold=True,
            size=DocumentFonts.EXCEL_TITLE_SIZE,
            color=DocumentColors.WHITE,
            name=DocumentFonts.EXCEL_TITLE_FONT
        ),
        'fill': PatternFill(
            start_color=DocumentColors.NAVY_BLUE,
            end_color=DocumentColors.NAVY_BLUE,
            fill_type="solid"
        ),
        'alignment': Alignment(
            horizontal="left",
            vertical="center"
        ),
        'border': Border(
            left=Side(style='thin', color=DocumentColors.BLACK),
            right=Side(style='thin', color=DocumentColors.BLACK),
            top=Side(style='thin', color=DocumentColors.BLACK),
            bottom=Side(style='thin', color=DocumentColors.BLACK)
        )
    }


def get_excel_header_style() -> Dict[str, Any]:
    """Get Excel header row style (medium blue background, white bold text)"""
    return {
        'font': Font(
            bold=True,
            size=DocumentFonts.EXCEL_HEADER_SIZE,
            color=DocumentColors.WHITE,
            name=DocumentFonts.EXCEL_HEADER_FONT
        ),
        'fill': PatternFill(
            start_color=DocumentColors.MEDIUM_BLUE,
            end_color=DocumentColors.MEDIUM_BLUE,
            fill_type="solid"
        ),
        'alignment': Alignment(
            horizontal="center",
            vertical="center",
            wrap_text=True
        ),
        'border': Border(
            left=Side(style='thin', color=DocumentColors.BLACK),
            right=Side(style='thin', color=DocumentColors.BLACK),
            top=Side(style='thin', color=DocumentColors.BLACK),
            bottom=Side(style='thin', color=DocumentColors.BLACK)
        )
    }


def get_excel_total_style() -> Dict[str, Any]:
    """Get Excel total/subtotal row style (light blue background, navy text)"""
    return {
        'font': Font(
            bold=True,
            size=DocumentFonts.EXCEL_BODY_SIZE,
            color=DocumentColors.NAVY_BLUE,
            name=DocumentFonts.EXCEL_BODY_FONT
        ),
        'fill': PatternFill(
            start_color=DocumentColors.LIGHT_BLUE,
            end_color=DocumentColors.LIGHT_BLUE,
            fill_type="solid"
        ),
        'alignment': Alignment(
            horizontal="right",
            vertical="center"
        ),
        'border': Border(
            left=Side(style='thin', color=DocumentColors.BLACK),
            right=Side(style='thin', color=DocumentColors.BLACK),
            top=Side(style='thin', color=DocumentColors.BLACK),
            bottom=Side(style='thin', color=DocumentColors.BLACK)
        )
    }


def get_excel_data_style(is_alternate_row: bool = False) -> Dict[str, Any]:
    """Get Excel data cell style with optional alternating row color"""
    fill_color = DocumentColors.LIGHT_GRAY if is_alternate_row else DocumentColors.WHITE

    return {
        'font': Font(
            size=DocumentFonts.EXCEL_BODY_SIZE,
            name=DocumentFonts.EXCEL_BODY_FONT
        ),
        'fill': PatternFill(
            start_color=fill_color,
            end_color=fill_color,
            fill_type="solid"
        ),
        'alignment': Alignment(
            horizontal="left",
            vertical="center"
        ),
        'border': Border(
            left=Side(style='thin', color=DocumentColors.BLACK),
            right=Side(style='thin', color=DocumentColors.BLACK),
            top=Side(style='thin', color=DocumentColors.BLACK),
            bottom=Side(style='thin', color=DocumentColors.BLACK)
        )
    }


# ============================================================================
# WORD STYLING FUNCTIONS
# ============================================================================

def apply_word_title_style(paragraph, text: str = None):
    """Apply professional title style to Word paragraph"""
    if text:
        paragraph.text = text

    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = DocumentFonts.WORD_TITLE_SIZE
    run.font.bold = True
    run.font.color.rgb = DocumentColors.NAVY_BLUE_RGB
    run.font.name = DocumentFonts.WORD_TITLE_FONT

    # Add spacing
    paragraph.space_before = Pt(0)
    paragraph.space_after = Pt(12)


def apply_word_heading1_style(paragraph, text: str = None):
    """Apply Heading 1 style (major sections)"""
    if text:
        paragraph.text = text

    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = DocumentFonts.WORD_HEADING1_SIZE
    run.font.bold = True
    run.font.color.rgb = DocumentColors.NAVY_BLUE_RGB
    run.font.name = DocumentFonts.WORD_HEADING_FONT

    paragraph.space_before = Pt(18)
    paragraph.space_after = Pt(6)


def apply_word_heading2_style(paragraph, text: str = None):
    """Apply Heading 2 style (subsections)"""
    if text:
        paragraph.text = text

    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = DocumentFonts.WORD_HEADING2_SIZE
    run.font.bold = True
    run.font.color.rgb = DocumentColors.MEDIUM_BLUE_RGB
    run.font.name = DocumentFonts.WORD_HEADING_FONT

    paragraph.space_before = Pt(12)
    paragraph.space_after = Pt(4)


def apply_word_body_style(paragraph, text: str = None):
    """Apply body text style"""
    if text:
        paragraph.text = text

    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.size = DocumentFonts.WORD_BODY_SIZE
    run.font.name = DocumentFonts.WORD_BODY_FONT

    paragraph.space_after = Pt(6)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY


def get_word_table_style_config() -> Dict[str, Any]:
    """Get configuration for Word table styling"""
    return {
        'header_fill': DocumentColors.MEDIUM_BLUE_RGB,
        'header_text': DocumentColors.WHITE_RGB,
        'alt_row_fill': DocumentColors.LIGHT_BLUE_RGB,
        'total_fill': DocumentColors.NAVY_BLUE_RGB,
        'total_text': DocumentColors.WHITE_RGB,
        'font_name': DocumentFonts.WORD_TABLE_FONT,
        'font_size': DocumentFonts.WORD_TABLE_SIZE,
        'border_color': DocumentColors.MEDIUM_GRAY
    }


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color to RGB for Word documents"""
    hex_color = hex_color.lstrip('#')
    r, g, b = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
    return RGBColor(r, g, b)


def get_status_color(value: float, threshold_good: float, threshold_bad: float) -> str:
    """Get color based on value thresholds"""
    if value >= threshold_good:
        return DocumentColors.SUCCESS_GREEN
    elif value <= threshold_bad:
        return DocumentColors.ERROR_RED
    else:
        return DocumentColors.WARNING_ORANGE


__all__ = [
    'DocumentColors',
    'DocumentFonts',
    'get_excel_title_style',
    'get_excel_header_style',
    'get_excel_total_style',
    'get_excel_data_style',
    'apply_word_title_style',
    'apply_word_heading1_style',
    'apply_word_heading2_style',
    'apply_word_body_style',
    'get_word_table_style_config',
    'hex_to_rgb',
    'get_status_color',
]
